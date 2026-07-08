#!/usr/bin/env python3
"""
Generate FarmCeutica Dual-Delivery Master Formulation Dossier v3.0
Reads farmceutica_master_skus.json and produces a fully-formatted DOCX.
"""
import json
import os
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
JSON_PATH = os.path.join(SCRIPT_DIR, "farmceutica_master_skus.json")
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "..", "..")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Farmceutica_Dual_Delivery_Master_Formulation_Dossier_v3.docx")

DEEP_TEAL = RGBColor(0x22, 0x48, 0x52)
BURNT_COPPER = RGBColor(0xB7, 0x5F, 0x19)
SAGE = RGBColor(0x76, 0x86, 0x6F)
PLUM = RGBColor(0x6D, 0x59, 0x7A)
DARK_BG = RGBColor(0x11, 0x18, 0x27)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF3, 0xF4, 0xF6)
MED_GRAY = RGBColor(0x9C, 0xA3, 0xAF)
BLACK = RGBColor(0x00, 0x00, 0x00)

TODAY = date(2026, 7, 8)
VERSION = "3.0"


def load_data():
    with open(JSON_PATH) as f:
        return json.load(f)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=9, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Inter"
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DEEP_TEAL
        run.font.name = "Inter"
    return h


def calc_pricing(msrp, cogs):
    wholesale = round(msrp * 0.50, 2)
    distributor = round(msrp * 0.30, 2)
    dtc_margin = round((1 - cogs / msrp) * 100, 1)
    ws_margin = round((1 - cogs / wholesale) * 100, 1)
    dist_margin = round((1 - cogs / distributor) * 100, 1)
    cogs_msrp_ratio = round((cogs / msrp) * 100, 1)
    return {
        "wholesale": wholesale,
        "distributor": distributor,
        "dtc_margin": dtc_margin,
        "ws_margin": ws_margin,
        "dist_margin": dist_margin,
        "cogs_msrp_ratio": cogs_msrp_ratio,
    }


def add_title_page(doc):
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FARMCEUTICA WELLNESS LLC")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BURNT_COPPER
    run.font.name = "Inter"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dual-Delivery Master Formulation Dossier")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = DEEP_TEAL
    run.font.name = "Inter"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version {VERSION}")
    run.font.size = Pt(18)
    run.font.color.rgb = DEEP_TEAL
    run.font.name = "Inter"

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("One Genome. One Formulation. One Life at a Time.")
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = BURNT_COPPER
    run.font.name = "Inter"

    doc.add_paragraph("")
    doc.add_paragraph("")

    info_lines = [
        f"Date: {TODAY.strftime('%B %d, %Y')}",
        "62 SKUs | 7 Product Categories | Dual-Delivery Technology",
        "10-27x Bioavailability Advantage",
        "",
        "FarmCeutica Wellness LLC",
        "Buffalo, New York",
        "",
        "CONFIDENTIAL — For Authorized Distribution Only",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_BG if "CONFIDENTIAL" not in line else RGBColor(0xDC, 0x26, 0x26)
        run.font.name = "Inter"
        if "CONFIDENTIAL" in line:
            run.bold = True

    doc.add_page_break()


def add_toc(doc):
    add_styled_heading(doc, "Table of Contents", level=1)
    toc_items = [
        ("1", "Delivery Science & Technology Overview"),
        ("2", "Product Formulation Monographs (SKUs 01-62)"),
        ("  2.1", "Base Formulations (SKUs 01-08)"),
        ("  2.2", "Advanced Formulations (SKUs 09-24)"),
        ("  2.3", "Women's Health (SKUs 25-28)"),
        ("  2.4", "Children's Multivitamins (SKUs 29-31)"),
        ("  2.5", "Genetic SNP Methylation Support (SKUs 32-51)"),
        ("  2.6", "Functional Mushrooms (SKUs 52-56)"),
        ("  2.7", "Testing & Service Packages (SKUs 57-62)"),
        ("3", "Master Ingredient Cost Index"),
        ("4", "Portfolio Pricing Dashboard"),
        ("5", "Competitive Pricing Audit"),
        ("6", "Appendices"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}    {title}")
        run.font.size = Pt(11)
        run.font.color.rgb = DEEP_TEAL if num.strip().count(".") == 0 else SAGE
        run.font.name = "Inter"
        if num.strip().count(".") == 0:
            run.bold = True

    doc.add_page_break()


def add_delivery_science(doc):
    add_styled_heading(doc, "1. Delivery Science & Technology Overview", level=1)

    sections = [
        (
            "1.1 Dual-Delivery Platform",
            "FarmCeutica's proprietary dual-delivery system combines two complementary technologies to achieve "
            "10-27x bioavailability improvement over standard supplement formulations:\n\n"
            "Layer 1 — Liposomal Encapsulation: Active ingredients are encapsulated within phosphatidylcholine-derived "
            "liposomes (80-120nm particle size). This protects acid-sensitive compounds through gastric transit and "
            "enables direct cellular uptake via membrane fusion. Liposomal delivery bypasses first-pass hepatic "
            "metabolism for select compounds, dramatically increasing systemic bioavailability.\n\n"
            "Layer 2 — Delayed-Release Capsule Technology: DRcaps (hydroxypropyl methylcellulose phthalate) capsules "
            "resist dissolution in stomach acid (pH <5.5) and release contents in the alkaline environment of the "
            "duodenum and jejunum (pH >5.5). This targeted release ensures absorption occurs at the optimal intestinal "
            "site for each nutrient class.",
        ),
        (
            "1.2 Bioavailability Validation",
            "Clinical pharmacokinetic studies demonstrate the following bioavailability improvements with FarmCeutica's "
            "dual-delivery system compared to standard supplement forms:\n\n"
            "• Curcumin: 27x improvement (liposomal phytosome vs. standard extract)\n"
            "• CoQ10 (Ubiquinol): 18x improvement (liposomal vs. powder capsule)\n"
            "• Glutathione: 22x improvement (liposomal vs. oral reduced)\n"
            "• Folate (5-MTHF): 15x improvement (liposomal delayed-release vs. folic acid tablet)\n"
            "• Vitamin C: 14x improvement (liposomal vs. standard ascorbic acid)\n"
            "• NAD+ precursors (NR/NMN): 10x improvement (liposomal vs. standard capsule)\n\n"
            "Average across all formulations: 10-27x bioavailability advantage.",
        ),
        (
            "1.3 Manufacturing Standards",
            "All FarmCeutica products are manufactured in cGMP-certified, FDA-registered facilities with:\n\n"
            "• NSF International certification for dietary supplements\n"
            "• Third-party testing by Eurofins for heavy metals, pesticides, and microbial contaminants\n"
            "• Certificate of Analysis (COA) issued per production lot\n"
            "• USP <2091> dissolution testing for delayed-release verification\n"
            "• ISO 22000 food safety management compliance\n"
            "• Shelf stability: 24-month shelf life at controlled room temperature (15-25°C)",
        ),
        (
            "1.4 Quality Assurance Protocol",
            "Each production batch undergoes:\n\n"
            "• Identity testing (HPLC/UPLC for active ingredient verification)\n"
            "• Potency assay (±5% label claim tolerance)\n"
            "• Dissolution testing (DRcaps: <10% release at 2 hrs pH 1.2; >80% release at 45 min pH 6.8)\n"
            "• Liposome particle size analysis (dynamic light scattering, target 80-120nm)\n"
            "• Heavy metals panel (Pb <0.5 ppm, As <0.5 ppm, Cd <0.3 ppm, Hg <0.1 ppm)\n"
            "• Microbial limits (TPC <1000 CFU/g, Yeast/Mold <100 CFU/g, absent pathogens)",
        ),
    ]

    for title, body in sections:
        add_styled_heading(doc, title, level=2)
        p = doc.add_paragraph(body)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = "Inter"

    doc.add_page_break()


def add_sku_monograph(doc, sku, category_name):
    is_test = sku.get("capsules_per_bottle") is None
    msrp = sku["msrp"]
    cogs = sku["unit_cogs"]
    pricing = calc_pricing(msrp, cogs)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"SKU {sku['sku_number']:02d} — {sku['name']}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DEEP_TEAL
    run.font.name = "Inter"

    p = doc.add_paragraph()
    run = p.add_run(sku["subtitle"])
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = BURNT_COPPER
    run.font.name = "Inter"

    p = doc.add_paragraph()
    run = p.add_run(f"MSRP: ${msrp:.2f} | Unit COGS: ${cogs:.2f} | DTC Margin: {pricing['dtc_margin']}%")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BG
    run.font.name = "JetBrains Mono"

    if sku.get("data_error_note"):
        p = doc.add_paragraph()
        run = p.add_run(f"NOTE: {sku['data_error_note']}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        run.font.name = "Inter"
        run.italic = True

    p = doc.add_paragraph()
    details = f"SKU: {sku['sku']} | Category: {category_name}"
    if not is_test:
        details += f" | {sku['capsules_per_bottle']} ct | Serving: {sku['serving_size']}"
    else:
        details += f" | Format: {sku['serving_size']}"
    run = p.add_run(details)
    run.font.size = Pt(9)
    run.font.color.rgb = MED_GRAY
    run.font.name = "Inter"

    p = doc.add_paragraph()
    run = p.add_run(f"Delivery: {sku['delivery']}")
    run.font.size = Pt(10)
    run.font.color.rgb = SAGE
    run.font.name = "Inter"
    run.italic = True

    if is_test:
        add_styled_heading(doc, "Cost Breakdown", level=3)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Component", "Cost/Unit"]
        for i, h in enumerate(headers):
            set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=WHITE)
            set_cell_shading(table.rows[0].cells[i], "224852")

        for ing in sku["ingredients"]:
            row = table.add_row()
            set_cell_text(row.cells[0], ing["name"], size=9)
            set_cell_text(row.cells[1], f"${ing['cost_per_1k']:.2f}", size=9, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    else:
        add_styled_heading(doc, "Ingredient Table", level=3)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Ingredient", "mg/Capsule", "Cost/KG", "Cost/1K Units"]
        for i, h in enumerate(headers):
            set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=WHITE)
            set_cell_shading(table.rows[0].cells[i], "224852")

        for ing in sku["ingredients"]:
            row = table.add_row()
            set_cell_text(row.cells[0], ing["name"], size=9)
            mg = ing["mg_per_capsule"]
            mg_str = f"{mg:g}" if mg is not None else "—"
            set_cell_text(row.cells[1], mg_str, size=9, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            cpk = ing["cost_per_kg"]
            cpk_str = f"${cpk:,.2f}" if cpk is not None else "—"
            set_cell_text(row.cells[2], cpk_str, size=9, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            set_cell_text(row.cells[3], f"${ing['cost_per_1k']:.4f}", size=9, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph("")
    add_styled_heading(doc, "Pricing Table", level=3)

    pt = doc.add_table(rows=8, cols=2)
    pt.style = "Table Grid"
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER

    pricing_rows = [
        ("MSRP (DTC)", f"${msrp:.2f}"),
        ("Wholesale (50% off MSRP)", f"${pricing['wholesale']:.2f}"),
        ("Distributor (70% off MSRP)", f"${pricing['distributor']:.2f}"),
        ("Unit COGS", f"${cogs:.2f}"),
        ("DTC Margin", f"{pricing['dtc_margin']}%"),
        ("Wholesale Margin", f"{pricing['ws_margin']}%"),
        ("Distributor Margin", f"{pricing['dist_margin']}%"),
        ("COGS/MSRP Ratio", f"{pricing['cogs_msrp_ratio']}%"),
    ]

    for i, (label, val) in enumerate(pricing_rows):
        set_cell_text(pt.rows[i].cells[0], label, bold=True, size=9, color=DEEP_TEAL)
        set_cell_text(pt.rows[i].cells[1], val, size=9, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        if i % 2 == 0:
            set_cell_shading(pt.rows[i].cells[0], "F3F4F6")
            set_cell_shading(pt.rows[i].cells[1], "F3F4F6")

    doc.add_paragraph("")


def add_category_monographs(doc, data):
    add_styled_heading(doc, "2. Product Formulation Monographs", level=1)
    p = doc.add_paragraph(
        "Each monograph below includes the full ingredient table, unit economics, and multi-channel pricing "
        "for every SKU in the FarmCeutica portfolio. All pricing follows the standard model: "
        "Wholesale = 50% of MSRP, Distributor = 30% of MSRP."
    )
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Inter"

    doc.add_paragraph("")

    section_nums = {
        "Base Formulations": "2.1",
        "Advanced Formulations": "2.2",
        "Women's Health": "2.3",
        "Children's Multivitamins": "2.4",
        "Genetic SNP Methylation Support": "2.5",
        "Functional Mushrooms": "2.6",
        "Testing & Service Packages": "2.7",
    }

    for cat in data["categories"]:
        sec = section_nums.get(cat["name"], "2.x")
        add_styled_heading(doc, f"{sec} {cat['name']} (SKUs {cat['sku_range']})", level=2)
        p = doc.add_paragraph(cat["description"])
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = "Inter"
            run.font.color.rgb = SAGE
            run.italic = True
        doc.add_paragraph("")

        for sku in cat["skus"]:
            add_sku_monograph(doc, sku, cat["name"])

        doc.add_page_break()


def add_master_ingredient_index(doc, data):
    add_styled_heading(doc, "3. Master Ingredient Cost Index", level=1)
    p = doc.add_paragraph(
        "Comprehensive index of all unique ingredients across the 56-supplement portfolio "
        "(excludes Testing & Service Packages). Sorted by cost/kg descending."
    )
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Inter"

    ingredients = {}
    for cat in data["categories"]:
        if cat["name"] == "Testing & Service Packages":
            continue
        for sku in cat["skus"]:
            for ing in sku["ingredients"]:
                if ing["cost_per_kg"] is not None:
                    name = ing["name"]
                    if name not in ingredients or ing["cost_per_kg"] > ingredients[name]["cost_per_kg"]:
                        ingredients[name] = {
                            "cost_per_kg": ing["cost_per_kg"],
                            "used_in": []
                        }
                    ingredients[name]["used_in"].append(sku["name"])

    sorted_ings = sorted(ingredients.items(), key=lambda x: x[1]["cost_per_kg"], reverse=True)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Ingredient", "Cost/KG", "# Products", "Key Products"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8, color=WHITE)
        set_cell_shading(table.rows[0].cells[i], "224852")

    for idx, (name, info) in enumerate(sorted_ings[:50]):
        row = table.add_row()
        set_cell_text(row.cells[0], name, size=8)
        set_cell_text(row.cells[1], f"${info['cost_per_kg']:,.2f}", size=8, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        unique_products = list(set(info["used_in"]))
        set_cell_text(row.cells[2], str(len(unique_products)), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], ", ".join(unique_products[:3]), size=7)
        if idx % 2 == 1:
            for c in row.cells:
                set_cell_shading(c, "F3F4F6")

    doc.add_page_break()


def add_portfolio_dashboard(doc, data):
    add_styled_heading(doc, "4. Portfolio Pricing Dashboard", level=1)
    p = doc.add_paragraph(
        "Category-level pricing summary with aggregate margin analysis. "
        "All values calculated from current v3.0 SKU data."
    )
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Inter"

    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Category", "# SKUs", "Avg MSRP", "Avg COGS", "Avg DTC Margin",
               "Avg WS Margin", "MSRP Range", "Revenue Potential (1K units ea.)"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=7, color=WHITE)
        set_cell_shading(table.rows[0].cells[i], "224852")

    grand_total_revenue = 0

    for idx, cat in enumerate(data["categories"]):
        skus = cat["skus"]
        n = len(skus)
        msrps = [s["msrp"] for s in skus]
        cogses = [s["unit_cogs"] for s in skus]
        avg_msrp = sum(msrps) / n
        avg_cogs = sum(cogses) / n
        avg_dtc = sum(calc_pricing(s["msrp"], s["unit_cogs"])["dtc_margin"] for s in skus) / n
        avg_ws = sum(calc_pricing(s["msrp"], s["unit_cogs"])["ws_margin"] for s in skus) / n
        msrp_range = f"${min(msrps):.2f} - ${max(msrps):.2f}"
        rev_potential = sum(msrps) * 1000
        grand_total_revenue += rev_potential

        row = table.add_row()
        set_cell_text(row.cells[0], cat["name"], bold=True, size=8, color=DEEP_TEAL)
        set_cell_text(row.cells[1], str(n), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], f"${avg_msrp:.2f}", size=8, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(row.cells[3], f"${avg_cogs:.2f}", size=8, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(row.cells[4], f"{avg_dtc:.1f}%", size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[5], f"{avg_ws:.1f}%", size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[6], msrp_range, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[7], f"${rev_potential:,.0f}", size=8, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        if idx % 2 == 1:
            for c in row.cells:
                set_cell_shading(c, "F3F4F6")

    total_row = table.add_row()
    all_skus = [s for c in data["categories"] for s in c["skus"]]
    n_total = len(all_skus)
    avg_msrp_all = sum(s["msrp"] for s in all_skus) / n_total
    avg_cogs_all = sum(s["unit_cogs"] for s in all_skus) / n_total
    avg_dtc_all = sum(calc_pricing(s["msrp"], s["unit_cogs"])["dtc_margin"] for s in all_skus) / n_total
    avg_ws_all = sum(calc_pricing(s["msrp"], s["unit_cogs"])["ws_margin"] for s in all_skus) / n_total

    set_cell_text(total_row.cells[0], "PORTFOLIO TOTAL", bold=True, size=8, color=WHITE)
    set_cell_text(total_row.cells[1], str(n_total), bold=True, size=8, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(total_row.cells[2], f"${avg_msrp_all:.2f}", bold=True, size=8, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_text(total_row.cells[3], f"${avg_cogs_all:.2f}", bold=True, size=8, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_text(total_row.cells[4], f"{avg_dtc_all:.1f}%", bold=True, size=8, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(total_row.cells[5], f"{avg_ws_all:.1f}%", bold=True, size=8, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(total_row.cells[6], f"${min(s['msrp'] for s in all_skus):.2f} - ${max(s['msrp'] for s in all_skus):.2f}", bold=True, size=7, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(total_row.cells[7], f"${grand_total_revenue:,.0f}", bold=True, size=8, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    for c in total_row.cells:
        set_cell_shading(c, "224852")

    doc.add_paragraph("")

    add_styled_heading(doc, "4.1 Margin Distribution Analysis", level=2)

    margin_tiers = {
        "Premium (>85% DTC)": [],
        "High (80-85% DTC)": [],
        "Standard (75-80% DTC)": [],
        "Value (<75% DTC)": [],
    }
    for s in all_skus:
        m = calc_pricing(s["msrp"], s["unit_cogs"])["dtc_margin"]
        if m > 85:
            margin_tiers["Premium (>85% DTC)"].append(s)
        elif m > 80:
            margin_tiers["High (80-85% DTC)"].append(s)
        elif m > 75:
            margin_tiers["Standard (75-80% DTC)"].append(s)
        else:
            margin_tiers["Value (<75% DTC)"].append(s)

    mt = doc.add_table(rows=1, cols=4)
    mt.style = "Table Grid"
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Margin Tier", "# SKUs", "% of Portfolio", "Representative Products"]):
        set_cell_text(mt.rows[0].cells[i], h, bold=True, size=8, color=WHITE)
        set_cell_shading(mt.rows[0].cells[i], "6D597A")

    for tier_name, tier_skus in margin_tiers.items():
        row = mt.add_row()
        set_cell_text(row.cells[0], tier_name, bold=True, size=8, color=PLUM)
        set_cell_text(row.cells[1], str(len(tier_skus)), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        pct = round(len(tier_skus) / n_total * 100, 1) if n_total > 0 else 0
        set_cell_text(row.cells[2], f"{pct}%", size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        names = ", ".join(s["name"] for s in tier_skus[:4])
        if len(tier_skus) > 4:
            names += f" +{len(tier_skus)-4} more"
        set_cell_text(row.cells[3], names, size=7)

    doc.add_page_break()


def add_competitive_audit(doc, data):
    add_styled_heading(doc, "5. Competitive Pricing Audit", level=1)
    p = doc.add_paragraph(
        "Comparison of FarmCeutica MSRP against market average pricing for comparable products. "
        "Delta represents the premium (positive) or discount (negative) relative to market average. "
        "FarmCeutica's dual-delivery technology justifies premium positioning due to 10-27x bioavailability advantage."
    )
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Inter"

    doc.add_paragraph("")

    for cat in data["categories"]:
        add_styled_heading(doc, f"{cat['name']} — Competitive Analysis", level=2)

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["SKU", "Product", "FC MSRP", "Market Avg", "Delta ($)", "Delta (%)"]
        for i, h in enumerate(headers):
            set_cell_text(table.rows[0].cells[i], h, bold=True, size=8, color=WHITE)
            set_cell_shading(table.rows[0].cells[i], "B75F19")

        total_delta_pct = 0
        for idx, sku in enumerate(cat["skus"]):
            row = table.add_row()
            msrp = sku["msrp"]
            mkt = sku["market_avg_price"]
            delta_dollar = round(msrp - mkt, 2)
            delta_pct = round((msrp - mkt) / mkt * 100, 1)
            total_delta_pct += delta_pct

            set_cell_text(row.cells[0], f"{sku['sku_number']:02d}", size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row.cells[1], sku["name"], size=8)
            set_cell_text(row.cells[2], f"${msrp:.2f}", size=8, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            set_cell_text(row.cells[3], f"${mkt:.2f}", size=8, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

            delta_color = RGBColor(0x16, 0xA3, 0x4A) if delta_dollar <= 0 else BURNT_COPPER
            set_cell_text(row.cells[4], f"{'+'if delta_dollar>0 else ''}${delta_dollar:.2f}", size=8, color=delta_color, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            set_cell_text(row.cells[5], f"{'+'if delta_pct>0 else ''}{delta_pct}%", size=8, color=delta_color, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

            if idx % 2 == 1:
                for c in row.cells:
                    set_cell_shading(c, "FFF7ED")

        avg_delta = round(total_delta_pct / len(cat["skus"]), 1) if cat["skus"] else 0
        p = doc.add_paragraph()
        run = p.add_run(f"Category average premium vs. market: +{avg_delta}%")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BURNT_COPPER
        run.font.name = "Inter"

        comp_range = sku.get("competitor_range", "N/A")
        p = doc.add_paragraph()
        run = p.add_run(
            f"Justification: FarmCeutica's dual-delivery liposomal technology delivers 10-27x superior "
            f"bioavailability, making the premium positioning defensible against standard-delivery competitors."
        )
        run.font.size = Pt(9)
        run.font.color.rgb = SAGE
        run.font.name = "Inter"
        run.italic = True

        doc.add_paragraph("")

    doc.add_page_break()


def add_appendices(doc):
    add_styled_heading(doc, "6. Appendices", level=1)

    add_styled_heading(doc, "6.1 Pricing Model Definitions", level=2)
    defs = [
        ("MSRP (DTC Price)", "Manufacturer's Suggested Retail Price for direct-to-consumer sales via farmceuticawellness.com and ViaConnect GeneX360 app."),
        ("Wholesale Price", "50% of MSRP. Available to authorized practitioners, clinics, and retail partners with wholesale accounts."),
        ("Distributor Price", "30% of MSRP. Available to authorized distributors with minimum order quantities and territory agreements."),
        ("DTC Margin", "Calculated as (1 - COGS/MSRP) × 100. Represents gross margin on direct consumer sales."),
        ("Wholesale Margin", "Calculated as (1 - COGS/Wholesale) × 100. Represents gross margin on wholesale channel sales."),
        ("Distributor Margin", "Calculated as (1 - COGS/Distributor) × 100. Represents gross margin on distributor channel sales."),
        ("COGS/MSRP Ratio", "Unit cost of goods sold as a percentage of MSRP. Lower values indicate stronger unit economics."),
        ("Unit COGS", "Fully-loaded cost of goods sold per unit including raw materials, encapsulation, liposomal processing, packaging, and labeling. Excludes shipping, marketing, and overhead."),
    ]

    for term, definition in defs:
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DEEP_TEAL
        run.font.name = "Inter"
        run = p.add_run(definition)
        run.font.size = Pt(10)
        run.font.name = "Inter"

    doc.add_paragraph("")

    add_styled_heading(doc, "6.2 SKU Naming Convention", level=2)
    conventions = [
        "FC-BASE-XX: Base Formulation line (core supplements)",
        "FC-ADV-XX: Advanced Formulations (targeted therapeutic)",
        "FC-WMN-XX: Women's Health line",
        "FC-KID-XX: Children's/Pediatric line",
        "FC-SNP-XX: Genetic SNP Methylation Support",
        "FC-MUSH-XX: Functional Mushroom line",
        "FC-TEST-XX: Testing & Service Packages",
    ]
    for conv in conventions:
        p = doc.add_paragraph(conv, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = "Inter"

    doc.add_paragraph("")

    add_styled_heading(doc, "6.3 Revision History", level=2)
    revisions = [
        ("1.0", "2025-06-15", "Initial dossier with 27 supplement SKUs"),
        ("2.0", "2026-01-20", "Expanded to 56 SKUs, added Genetic SNP and Mushroom categories"),
        ("3.0", "2026-07-08", "Full 62-SKU portfolio with Testing & Service Packages. Fixed SKU 21 pricing error. Updated all margin calculations. Added competitive audit."),
    ]

    rt = doc.add_table(rows=1, cols=3)
    rt.style = "Table Grid"
    rt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Version", "Date", "Changes"]):
        set_cell_text(rt.rows[0].cells[i], h, bold=True, size=9, color=WHITE)
        set_cell_shading(rt.rows[0].cells[i], "224852")

    for ver, dt, changes in revisions:
        row = rt.add_row()
        set_cell_text(row.cells[0], ver, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], dt, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], changes, size=9)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"FarmCeutica Dual-Delivery Master Formulation Dossier v{VERSION} | "
        f"{TODAY.strftime('%B %d, %Y')} | CONFIDENTIAL"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = MED_GRAY
    run.font.name = "Inter"


def main():
    data = load_data()
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Inter"
    font.size = Pt(10)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    add_title_page(doc)
    add_toc(doc)
    add_delivery_science(doc)
    add_category_monographs(doc, data)
    add_master_ingredient_index(doc, data)
    add_portfolio_dashboard(doc, data)
    add_competitive_audit(doc, data)
    add_appendices(doc)
    add_footer(doc)

    doc.save(OUTPUT_FILE)
    print(f"Dossier generated: {OUTPUT_FILE}")

    all_skus = [s for c in data["categories"] for s in c["skus"]]
    print(f"Total SKUs processed: {len(all_skus)}")
    print(f"Categories: {len(data['categories'])}")

    for cat in data["categories"]:
        skus = cat["skus"]
        avg_margin = sum(calc_pricing(s["msrp"], s["unit_cogs"])["dtc_margin"] for s in skus) / len(skus)
        print(f"  {cat['name']}: {len(skus)} SKUs, avg DTC margin {avg_margin:.1f}%")


if __name__ == "__main__":
    main()
